"""Упрощённое преобразование Markdown → DOCX (python-docx).

Поддерживаются: заголовки ``#``–``###``, маркированные списки ``-``/``*``, абзацы,
жирный ``**…**``, курсив ``*…*``, блоки кода `` ``` `` (моноширинный текст)."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt


def _append_inline_runs(paragraph, text: str, *, monospace: bool = False) -> None:
    """Разбор ``**жирный**`` и ``*курсив*`` (одиночная звёздочка, не пара)."""
    pos = 0
    n = len(text)
    while pos < n:
        if text.startswith("**", pos):
            end = text.find("**", pos + 2)
            if end == -1:
                run = paragraph.add_run(text[pos:])
                run.bold = True
                if monospace:
                    run.font.name = "Courier New"
                break
            run = paragraph.add_run(text[pos + 2 : end])
            run.bold = True
            if monospace:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            pos = end + 2
            continue
        if text[pos] == "*" and (pos + 1 >= n or text[pos + 1] != "*"):
            end = text.find("*", pos + 1)
            if end == -1:
                paragraph.add_run(text[pos:])
                break
            run = paragraph.add_run(text[pos + 1 : end])
            run.italic = True
            if monospace:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            pos = end + 1
            continue
        next_star = text.find("*", pos + 1)
        next_bold = text.find("**", pos + 2)
        candidates = [x for x in (next_star, next_bold) if x != -1]
        nxt = min(candidates) if candidates else -1
        if nxt == -1:
            run = paragraph.add_run(text[pos:])
            if monospace:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            break
        if nxt > pos:
            run = paragraph.add_run(text[pos:nxt])
            if monospace:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
        pos = nxt


def markdown_to_docx_bytes(markdown: str, *, title: str | None = None) -> bytes:
    """Возвращает байты ``.docx`` (Office Open XML)."""

    doc = Document()
    if title and title.strip():
        doc.add_heading(title.strip(), level=1)

    text = (markdown or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines, in_code
        if not code_lines:
            in_code = False
            return
        p = doc.add_paragraph()
        _append_inline_runs(p, "\n".join(code_lines), monospace=True)
        code_lines = []
        in_code = False

    for raw in lines:
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(raw)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        hm = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if hm:
            depth = len(hm.group(1))
            rest = hm.group(2).strip()
            word_level = min(max(depth, 1), 3)
            doc.add_heading(rest, level=word_level)
            continue

        if re.match(r"^[-*]\s+", stripped):
            body = re.sub(r"^[-*]\s+", "", stripped)
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
                body = f"• {body}"
            _append_inline_runs(p, body)
            continue

        p = doc.add_paragraph()
        _append_inline_runs(p, stripped)

    if in_code and code_lines:
        flush_code()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
